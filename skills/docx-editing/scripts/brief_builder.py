#!/usr/bin/env python3
"""Kang Lab Research Brief builder — SHARED python-docx helpers (canonical).

Usage from a project's create_research_brief.py:
    import sys; sys.path.insert(0, r"D:\\onedrive\\claude\\agent\\tools")
    from brief_builder import (new_doc, title_block, h1, h2, body, bullet, gap,
                               info_table, data_table, csv_table, figure, page_break)
    doc = new_doc()
    title_block(doc, "<paper title>", "Target: <journal>  ·  2026-07-12")
    h1(doc, "Study Overview"); info_table(doc, [...]); ...
    doc.save(path)

Import — DON'T copy per project (drift). This dir is OneDrive-synced across machines.

Encodes feedback/manuscript_rules.md "Research Brief" rules:
  - EN section headings (h1/h2) + KO body prose (body/bullet) + EN tables/figures (논문 동일본).
  - NO colored key-finding box — brief != PPT (AI 티 + 흐름 방해). Intentionally NOT provided here.
  - lean: core tables/figures only; Secondary/Sensitivity mostly prose.
  - figures embedded titleless (title/subtitle belong in the caption, not baked into the PNG).
  - Korean renders via Normal-style eastAsia = 맑은 고딕 (Calibri lacks Hangul glyphs).
  - info_table uses fixed layout + per-cell widths (Word ignores table.columns[i].width).
"""
import csv, os
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GRAY      = RGBColor(0x66, 0x66, 0x66)
KFONT     = "맑은 고딕"

# Table body size. Was 8.5 until 2026-08-19, when the user re-sized every table in
# a delivered brief by hand: 8.0/8.5/9.0 -> 10.0 across the board, with the prose
# left untouched. feedback/visual_qa.md already carried "BODY_PT <- 10" for DOCX
# tables, but only under a flextable heading, so it never reached this module and
# the too-small default kept winning. It is a default here so remembering is not
# required. Room for a bigger table comes from cutting columns/words, never from
# shrinking type. Captions/footnotes stay smaller (see note-style callers).
BODY_TABLE_PT = 10


def kfont(run):
    """Set East Asian font on a run so Hangul renders (Calibri has no Hangul glyphs)."""
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), KFONT)


def new_doc(margin_cm=2.3, body_pt=10):
    """Fresh Document with lab-standard Normal style (Calibri + 맑은 고딕) and margins."""
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(body_pt)
    st.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), KFONT)
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(margin_cm)
    sec.left_margin = sec.right_margin = Cm(margin_cm)
    return doc


def set_cell_shading(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    tcPr.append(tcPr.makeelement(qn("w:shd"),
        {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex}))


def style_header_row(row, bg="1F4E79", fs=9):
    for cell in row.cells:
        set_cell_shading(cell, bg)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = WHITE; run.font.size = Pt(fs); run.font.bold = True


def title_block(doc, title, meta, label="Research Brief (논문요약서)"):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label); r.font.size = Pt(13); r.bold = True; r.font.color.rgb = DARK_BLUE; kfont(r)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title); r.font.size = Pt(12.5); r.bold = True; kfont(r)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta); r.font.size = Pt(9); r.italic = True; r.font.color.rgb = GRAY; kfont(r)


def h1(doc, text, page_break_before=False):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
    if page_break_before:
        p.paragraph_format.page_break_before = True   # avoids stray blank page from a standalone add_page_break()
    r = p.add_run(text); r.font.size = Pt(14); r.font.color.rgb = DARK_BLUE; r.bold = True; kfont(r)
    pPr = p._element.get_or_add_pPr(); pBdr = pPr.makeelement(qn("w:pBdr"), {})
    pBdr.append(pBdr.makeelement(qn("w:bottom"),
        {qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "1", qn("w:color"): "1F4E79"}))
    pPr.append(pBdr)


def h2(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(11); r.font.color.rgb = DARK_BLUE; r.bold = True; kfont(r)


def body(doc, text, fs=9.5):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.size = Pt(fs); kfont(r); return p


def bullet(doc, text, prefix=None, fs=9.5):
    p = doc.add_paragraph(style="List Bullet")
    if prefix:
        rb = p.add_run(prefix); rb.bold = True; rb.font.size = Pt(fs); kfont(rb)
    r = p.add_run(text); r.font.size = Pt(fs); kfont(r); return p


def gap(doc, pt=2):
    doc.add_paragraph().paragraph_format.space_after = Pt(pt)


def page_break(doc):
    doc.add_page_break()


def info_table(doc, items, label_cm=3.3, value_cm=13.7):
    """Two-column key/value table. Fixed layout + per-cell widths (Word ignores columns[i].width)."""
    t = doc.add_table(rows=len(items), cols=2); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; t.allow_autofit = False
    tblPr = t._tbl.tblPr
    lay = tblPr.find(qn("w:tblLayout"))
    if lay is None:
        lay = tblPr.makeelement(qn("w:tblLayout"), {}); tblPr.append(lay)
    lay.set(qn("w:type"), "fixed")
    for i, (k, v) in enumerate(items):
        c0, c1 = t.rows[i].cells
        c0.width = Cm(label_cm); c1.width = Cm(value_cm)
        r0 = c0.paragraphs[0].add_run(k); r0.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = DARK_BLUE; kfont(r0)
        r1 = c1.paragraphs[0].add_run(v); r1.font.size = Pt(9); kfont(r1)
        set_cell_shading(c0, "E8F0F8")
        if i % 2 == 1: set_cell_shading(c1, "F8F8F8")
    return t


def _repeat_header_row(t):
    """Mark row 0 as a header row so Word repeats it when the table spans pages.

    Without this a long table breaks mid-row on page 2 with no header, and the reader
    cannot tell which column is which. Word only does this when the row is tagged, so
    it has to be code, not a rule someone remembers. (airkorea 2026-08-14: caught in
    render QA after the source table split across pages headerless.)
    """
    trPr = t.rows[0]._tr.get_or_add_trPr()
    el = trPr.makeelement(qn("w:tblHeader"), {qn("w:val"): "true"})
    trPr.append(el)
    return t


def data_table(doc, headers, rows, fs=BODY_TABLE_PT, keep_together=False, widths=None):
    """Table from in-memory headers + rows (navy header, zebra body).

    The header row repeats across page breaks by default.

    widths: per-column cm. Without it Word divides the width evenly, which wraps
    long label columns onto two lines while numeric columns sit half empty. Give
    the label column the room it needs and the number columns only what they use
    (list should sum to the text width, ~16.99cm at the default margins).
    """
    t = doc.add_table(rows=1 + len(rows), cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hd in enumerate(headers):
        rr = t.rows[0].cells[j].paragraphs[0].add_run(str(hd)); rr.font.size = Pt(fs); kfont(rr)
    style_header_row(t.rows[0], fs=fs)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            rr = cell.paragraphs[0].add_run(str(val)); rr.font.size = Pt(fs); kfont(rr)
            if i % 2 == 1: set_cell_shading(cell, "F2F2F2")
    if widths:
        _set_col_widths(t, widths)
    if keep_together:
        _keep_rows_together(t)
    return _repeat_header_row(t)


def csv_table(doc, path, fs=BODY_TABLE_PT, table1=False, widths=None, keep_together=False):
    """Table from a CSV file (never hand-typed). table1=True: bold category-header rows
    (empty data cells) and indent sub-items (leading spaces) per the Table 1 standard."""
    with open(path, newline="", encoding="utf-8-sig") as f:
        rows = list(csv.reader(f))
    if not table1:
        return data_table(doc, rows[0], rows[1:], fs=fs, keep_together=keep_together)
    hdr, brows = rows[0], rows[1:]
    t = doc.add_table(rows=1 + len(brows), cols=len(hdr)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(hdr):
        rr = t.rows[0].cells[j].paragraphs[0].add_run(h); rr.font.size = Pt(fs); kfont(rr)
    style_header_row(t.rows[0], fs=fs)
    for i, row in enumerate(brows):
        is_cat = all(c.strip() == "" for c in row[1:])
        indent = row[0].startswith(" ")
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]; p = cell.paragraphs[0]
            rr = p.add_run(val.strip()); rr.font.size = Pt(fs); kfont(rr)
            if j == 0 and is_cat: rr.bold = True
            if j == 0 and indent: p.paragraph_format.left_indent = Cm(0.35)
    if widths:
        _set_col_widths(t, widths)
    if keep_together:
        _keep_rows_together(t)
    return _repeat_header_row(t)


def _keep_rows_together(table):
    """Stop a short table splitting across a page break (which strands a lone row).
    keep_with_next on every row but the last -> Word moves the whole table to the next page.
    Do NOT use on long tables (e.g. Table 1) that legitimately must span pages."""
    rows = table.rows
    for r in rows[:-1]:
        for c in r.cells:
            for para in c.paragraphs:
                para.paragraph_format.keep_with_next = True
    return table

def _set_col_widths(table, widths_cm):
    """Word ignores columns[i].width; fix the layout and stamp width on every cell
    (same idiom as info_table).

    The tblGrid is updated as well. Word honours the per-cell widths under a fixed
    layout either way, but a stale grid makes the table LIE to anything that reads
    w:gridCol - including tools/audit_table_widths.py, which then measures the old
    even split and reports phantom (or missed) overflow. Two readers of the same
    table must not get different answers. (2026-08-19: cost a wrong "widths did not
    apply" diagnosis in both directions.)
    """
    table.autofit = False; table.allow_autofit = False
    tblPr = table._tbl.tblPr
    lay = tblPr.find(qn("w:tblLayout"))
    if lay is None:
        lay = tblPr.makeelement(qn("w:tblLayout"), {}); tblPr.append(lay)
    lay.set(qn("w:type"), "fixed")
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for i, gc in enumerate(grid.findall(qn("w:gridCol"))):
            if i < len(widths_cm):
                gc.set(qn("w:w"), str(int(Cm(widths_cm[i]).twips)))
    return table

def figure(doc, path, caption, max_w_cm=15.0, max_h_cm=21.0):
    """Embed a figure 1:1 (aspect preserved) + gray italic caption. Figures must be titleless
    (no baked-in title/subtitle) — that text belongs in `caption`."""
    path = str(path)
    if not os.path.exists(path):
        body(doc, f"[missing figure: {path}]"); return
    w, h = Image.open(path).size
    width = Cm(max_w_cm); height = Cm(max_w_cm * h / w)
    if height > Cm(max_h_cm):
        height = Cm(max_h_cm); width = Cm(max_h_cm * w / h)
    doc.add_picture(path, width=width)
    pic_par = doc.paragraphs[-1]
    pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Caption must never break away from its figure. Without this the caption can
    # land alone on the next page - or, worse, at the top of a page directly above
    # a DIFFERENT figure, where it reads as that figure's caption. (2026-08-19:
    # both happened at once in one brief, and it was worked around by hand-shrinking
    # the image until figure+caption fit together. keep_with_next fixes the cause.)
    pic_par.paragraph_format.keep_with_next = True
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(10)
    r = cp.add_run(caption); r.font.size = Pt(9); r.font.color.rgb = GRAY; r.italic = True; kfont(r)
