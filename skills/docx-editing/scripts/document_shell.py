#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""document_shell -- 장르의 «골격». 부품(문단·표·렌더) 위에 한 층 더.

    from docx_kit import manuscript_shell, brief_shell

왜 있나 (2026-08-25 실측)
────────────────────────
docx 를 만드는 스크립트가 64개인데 kit(부품)을 쓰는 것은 11개였다. 나머지 53개 중
**51개가 「굵게 처리」를 다시 짰다** -- kit 에 있는데도. 부품을 모은 것만으로는
안 멈춘 것이다.

이유는 부품이 모자라서가 아니라 **골격을 아무도 안 줬기 때문**이다. 「원고 totale」
하나를 짜려면 표지·쪽나눔·절 조립·표 번호·그림 캡션·게이트를 매번 새로 엮어야 했고,
그 과정에서 부품까지 다시 짜게 된다. 실측: 같은 골격의 totale 빌더가 **7개**
(254~741줄)로 흩어져 있었다.

여기 있는 것은 그 골격 둘이다 -- 각각 여러 빌더에서 «반복된 것»만 추렸다.

⚠️⚠️ **골격은 문장을 만들지 않는다.** 산문은 사람이 쓰고 골격은 배치·번호·게이트만
   맡는다. 그래야 재분석 뒤에 숫자만 따라오고 서사는 사람이 책임진다.

골격이 없는 장르 (정부·감독자 보고서, 수업자료)
──────────────────────────────────────────────
**측정하지 않았으므로 만들지 않았다.** 부품으로 짓는다. 같은 골격이 «셋 이상»
반복되는 것이 확인되면 그때 여기로 올린다 -- 그것이 이 파일이 생긴 근거다.
"""
import glob as _glob
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from manuscript_table import add_journal_table
from md_to_docx import render_markdown

__all__ = ["manuscript_shell", "brief_shell", "fill_values", "GateError"]

# ⚠️ [a-zA-Z0-9_]로 좁히지 말 것 -- 한글 키({{미검증_수}})가 조용히 통과해 그대로
#    원고에 인쇄된다. 게이트가 막으라는 바로 그 실패다(2026-08-25 셀프테스트가 잡음).
PLACEHOLDER = re.compile(r"\{\{\s*([^{}\n]+?)\s*\}\}")
REF_MARK = re.compile(r"\[REF:\s*([^\]]+)\]")


class GateError(RuntimeError):
    """게이트가 막았다. 빌드를 «멈춘다» -- 경고로 흘리면 아무도 안 본다."""


# ── 숫자 게이트 ────────────────────────────────────────────────────────────
def fill_values(text, values, where=""):
    """{{키}}를 values로 치환한다. 하나라도 못 채우면 GateError.

    ⚠️ 「없으면 빈 칸으로 두기」를 만들지 말 것 -- 숫자가 빠진 원고가 조용히 나간다.
    """
    missing = []

    def sub(m):
        k = m.group(1)
        if k not in values:
            missing.append(k)
            return m.group(0)
        return str(values[k])

    out = PLACEHOLDER.sub(sub, text)
    if missing:
        raise GateError(f"치환되지 않은 자리표시자{' (' + where + ')' if where else ''}: "
                        + ", ".join(sorted(set(missing))))
    return out


def _words(text):
    """[REF: ...] 표시와 자리표시자를 뺀 «본문» 단어 수."""
    text = REF_MARK.sub("", PLACEHOLDER.sub("", text))
    return len(text.split())


def _load_sections(sections, values):
    """sections: (제목, 마크다운) 목록 / .md 경로 목록 / 글롭 문자열.
    반환: [(stem, 제목 또는 None, 치환된 마크다운)]"""
    items = []
    if isinstance(sections, str):
        sections = sorted(_glob.glob(sections))
    for s in sections:
        if isinstance(s, (tuple, list)):
            title, md = s[0], s[1]
            stem = title or ""
        else:
            p = Path(s)
            stem, title, md = p.stem, None, p.read_text(encoding="utf-8")
        if values is not None:
            md = fill_values(md, values, where=stem)
        items.append((stem, title, md))
    return items


# ── 골격 1: 원고 totale ────────────────────────────────────────────────────
def manuscript_shell(out, title, sections, *, authors=None, abstract=None,
                     cover=None, tables=(), figures=(), values=None,
                     word_limit=None, no_count=(), subtitle=None):
    """원고 totale 을 조립한다. 검토용 표지는 «쪽을 끊어» 붙으므로 투고 시 통째로 뺀다.

    배치:  [표지] → 제목·저자 → [초록] → 본문 절 → [표] → [그림]

    sections   (제목, 마크다운) 목록 · .md 경로 목록 · 글롭 문자열
    cover      검토용 표지. (제목, 마크다운) 또는 마크다운 문자열. 쪽나눔으로 분리된다
    tables     [(캡션, rows[, col_widths])] -- rows[0]이 헤더. 저널 3선표(세로줄 없음).
               ⚠️ col_widths(인치)를 주는 것이 기본이다 -- 안 주면 Word가 자동배분해
               라벨 열이 두 줄로 접히고 숫자 열이 반쯤 비운다
    figures    [(캡션, 이미지경로, 폭인치)] -- 폭 생략 시 6.0
    values     {{키}} 사전. 주면 미치환 시 «멈춘다»
    word_limit 본문 단어 상한. 넘으면 «멈춘다»
    no_count   본문 단어에 세지 않을 절의 stem 들(선언·감사·참고문헌 등)

    반환: {"body_words", "per_section", "refs", "path"} -- 보고는 호출자가 한다.
    """
    items = _load_sections(sections, values)
    if abstract is not None and values is not None:
        abstract = fill_values(abstract, values, where="abstract")
    if isinstance(cover, (tuple, list)):
        cover = cover[1]
    if cover is not None and values is not None:
        cover = fill_values(cover, values, where="cover")

    per_section, refs, body_words = {}, [], 0
    for stem, _t, md in items:
        n = _words(md)
        per_section[stem] = n
        if stem not in no_count:
            body_words += n
        refs += [(stem, r) for r in REF_MARK.findall(md)]

    if word_limit and body_words > word_limit:
        raise GateError(f"본문이 {body_words - word_limit}단어 초과 "
                        f"({body_words}/{word_limit}). 절별: {per_section}")

    doc = Document()

    if cover:
        render_markdown(doc, cover)
        doc.add_page_break()

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(title)
    r.bold = True
    r.font.size = Pt(14)
    for line, size, italic in ((authors, 11, False), (subtitle, 9, True)):
        if line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = p.add_run(line)
            rr.font.size = Pt(size)
            rr.italic = italic

    if abstract:
        doc.add_paragraph()
        render_markdown(doc, abstract)
        doc.add_page_break()

    for i, (_stem, sec_title, md) in enumerate(items):
        if sec_title:
            doc.add_heading(sec_title, level=1)
        render_markdown(doc, md)

    if tables:
        doc.add_page_break()
        for i, tb in enumerate(tables, 1):
            caption, rows = tb[0], tb[1]
            widths = tb[2] if len(tb) > 2 else None
            cap = doc.add_paragraph()                      # 제목은 표 «위» (저널 관습)
            cr = cap.add_run(f"Table {i}. {caption}" if caption else f"Table {i}")
            cr.bold = True
            cr.font.size = Pt(10)
            cap.paragraph_format.keep_with_next = True     # 제목이 표와 떨어지지 않게
            add_journal_table(doc, rows, col_widths=widths)
            doc.add_paragraph()

    if figures:
        doc.add_page_break()
        for i, fig in enumerate(figures, 1):
            path, caption = fig[1], fig[0]
            width = fig[2] if len(fig) > 2 else 6.0
            doc.add_picture(str(path), width=Inches(width))
            pic = doc.paragraphs[-1]
            pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic.paragraph_format.keep_with_next = True     # 캡션이 다른 그림에 붙지 않게
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cap.add_run(f"Figure {i}. {caption}" if caption else f"Figure {i}")
            cr.font.size = Pt(9)

    Path(out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return {"body_words": body_words, "per_section": per_section,
            "refs": refs, "path": str(out)}


# ── 골격 2: 국문 브리프·사례보고서 ─────────────────────────────────────────
def brief_shell(out, title, blocks, *, subtitle=None, values=None, footer=None):
    """국문 브리프. 격자·색 헤더가 «의도된» 디자인인 장르다(저널 표를 쓰지 않는다).

    blocks: 순서대로 배치할 목록. 각 항목은
        ("h", "절 제목")            소제목
        ("md", "마크다운")          산문
        ("table", 캡션, rows)       격자표. rows[0]이 헤더 (저널 표와 같은 모양으로 준다)
        ("fig", 캡션, 경로[, 폭])   그림 + 캡션
        ("break", None)             쪽나눔
    """
    from brief_builder import data_table                    # 순환 import 회피

    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(16)
    if subtitle:
        sp = doc.add_paragraph()
        sr = sp.add_run(subtitle)
        sr.font.size = Pt(9)
        sr.italic = True
    doc.add_paragraph()

    for b in blocks:
        kind = b[0]
        if kind == "break":
            doc.add_page_break()
        elif kind == "h":
            doc.add_heading(b[1], level=1)
        elif kind == "md":
            md = b[1]
            if values is not None:
                md = fill_values(md, values, where="brief")
            render_markdown(doc, md)
        elif kind == "table":
            if b[1]:
                cap = doc.add_paragraph()
                cr = cap.add_run(b[1])
                cr.bold = True
                cr.font.size = Pt(10)
                cap.paragraph_format.keep_with_next = True
            rows = b[2]
            data_table(doc, rows[0], rows[1:])             # rows[0] = 헤더
            doc.add_paragraph()
        elif kind == "fig":
            width = b[3] if len(b) > 3 else 6.0
            doc.add_picture(str(b[2]), width=Inches(width))
            pic = doc.paragraphs[-1]
            pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic.paragraph_format.keep_with_next = True
            if b[1]:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cap.add_run(b[1])
                cr.font.size = Pt(9)
        else:
            raise ValueError(f"모르는 블록 종류: {kind}")

    if footer:
        doc.add_paragraph()
        fp = doc.add_paragraph()
        fr = fp.add_run(footer)
        fr.font.size = Pt(8)
        fr.italic = True

    Path(out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return {"path": str(out)}
