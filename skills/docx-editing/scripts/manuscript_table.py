"""Journal-style tables for python-docx manuscripts.

Import this instead of hand-rolling table code. Every rule the lab has been
given about manuscript tables is the default here, so a new pipeline gets them
without anyone having to remember them:

  * no vertical rules, ever; three horizontal rules (top / under-header / bottom)
    and none between body rows
  * a variable is a bold label row spanning the width; its categories sit under
    it, indented - never separated by a rule
  * a variable's p-value rides on its label row, not on its last category
  * cell padding and paragraph spacing set explicitly (Word's defaults make rows
    roughly half again too tall for a typeset table)
  * significance stars reduced to the levels the table actually uses
  * long tables reprint their header on continuation pages

Typical use:

    from manuscript_table import add_journal_table, significance_notation

    rows, legend = significance_notation(read_xlsx_rows(path))
    add_journal_table(doc, rows, col_widths=[2.15, 1.15, 1.15, 1.15, 0.7])

Row convention: `rows` is a list of lists of strings, `rows[0]` the header.
A body row whose data cells are all empty is a group header. A body row whose
first cell starts with two spaces is a category of the group above it.

R pipelines: the flextable equivalent is `theme_booktabs()` plus
`padding(padding = 2)`; do not use `set_flextable_defaults` themes that draw
vertical borders.

Extracted 2026-07-20 from a manuscript builder after four rounds of supervisor
review. Rules and rationale: references/docx-conventions.md.
"""
import re

from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

__all__ = [
    "add_journal_table", "significance_notation", "is_group_row",
    "pvalue_columns", "set_journal_borders", "set_cell_margins",
    "set_row_bottom_rule", "repeat_header_row", "strip_footnote_rows",
]

DEFAULT_FONT = "Times New Roman"
P_HEADER_RE = re.compile(r"^p[ -]?(value)?$", re.I)
# "1.632 (1.440-1.850)" and "1.00 (ref)" are estimates; "18,952", "3-4" and "26.9" are not.
ESTIMATE_RE = re.compile(r"^-?\d[\d,.]*\s*\(.+\)\s*\**$")


# ---------------------------------------------------------------------------
# row-level helpers (usable without a document, e.g. by a markdown builder)
# ---------------------------------------------------------------------------

def pvalue_columns(header):
    """Indices of p-value columns, which a group header row is allowed to fill.

    A variable's p-value belongs on its label row - the test is for the whole
    variable, not its last category - so that one filled cell must not
    disqualify the row from being recognised as a group header."""
    return {i for i, h in enumerate(header) if P_HEADER_RE.match(str(h).strip())}


def is_group_row(row, ignore_cols=()):
    """True if `row` carries a label and no data, ignoring `ignore_cols`."""
    data = [c for i, c in enumerate(row) if i > 0 and i not in ignore_cols]
    return bool(str(row[0]).strip()) and not any(str(c).strip() for c in data)


def significance_notation(rows):
    """Reduce significance marking to what the table actually needs.

    Returns `(rows, legend)`. With several levels present, stars stay and the
    legend lists only those levels. When ONE level is present AND every estimate
    in the table carries it, the stars add nothing a sentence cannot: they are
    stripped and the legend states the level for all of them.

    ⚠️ The universal claim is only made when it is true. A table where one level
    is present but some estimates are unmarked - a cause-specific table with a
    null cancer row, say - keeps its stars and gets `* p<0.001.` instead. The
    earlier version stripped on level-count alone and printed "All estimates
    p<0.001." underneath a confidence interval spanning 1.00 (2026-08-29).

    Note the `rstrip()` - leading spaces are a row's indent, not padding, and
    stripping them silently flattens the table's hierarchy."""
    levels = (("*", "p<0.05"), ("**", "p<0.01"), ("***", "p<0.001"))
    text = " ".join(str(c) for r in rows for c in r)
    present = [(sym, lvl) for sym, lvl in levels
               if re.search(r"(?<!\*)" + re.escape(sym) + r"(?!\*)", text)]
    if not present:
        return rows, ""
    if len(present) > 1:
        return rows, ", ".join(sym + lvl for sym, lvl in present) + "."
    sym, lvl = present[0]
    # An estimate is a cell carrying a number and a parenthesised interval; a count, an N or a
    # percentage is not one, and must not decide whether the stars can go.
    est = [str(c).rstrip() for r in rows[1:] for c in r
           if ESTIMATE_RE.match(str(c).strip())]
    if est and all(c.endswith(sym) for c in est):
        stripped = [[re.sub(r"\*+$", "", str(c).rstrip()) for c in r] for r in rows]
        return stripped, "All estimates %s." % lvl
    return rows, "%s %s." % (sym, lvl)


def strip_footnote_rows(rows):
    """Drop trailing footnote rows baked into a source sheet.

    A footnote row has sentence-length content in the first column only;
    rendered as a table row it crams the note into the leftmost cell. Pass the
    note to the builder's footnote argument instead."""
    while len(rows) > 1:
        col0 = str(rows[-1][0]).strip()
        if col0 and len(col0) > 40 and not any(str(c).strip() for c in rows[-1][1:]):
            rows = rows[:-1]
        else:
            break
    return rows


# ---------------------------------------------------------------------------
# docx-level helpers
# ---------------------------------------------------------------------------

def _border_el(tag, width=None):
    """One border element: a rule of `width` eighth-points, or explicitly none."""
    el = OxmlElement("w:%s" % tag)
    if width is None:
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
    else:
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(width))
        el.set(qn("w:color"), "000000")
    el.set(qn("w:space"), "0")
    return el


def set_journal_borders(table, width=12):
    """Top and bottom rules only - no verticals, no rule between body rows.

    python-docx's built-in `Table Grid` style draws all four edges plus both
    inside grids; assigning any built-in grid/accent style undoes this."""
    tblPr = table._tbl.tblPr
    for el in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(el)
    borders = OxmlElement("w:tblBorders")
    for tag, w in (("top", width), ("left", None), ("bottom", width),
                   ("right", None), ("insideH", None), ("insideV", None)):
        borders.append(_border_el(tag, w))     # schema order matters
    tblPr.append(borders)


def set_row_bottom_rule(row, width=6):
    """Draw a rule under one row - the under-header rule."""
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        for el in tcPr.findall(qn("w:tcBorders")):
            tcPr.remove(el)
        tcb = OxmlElement("w:tcBorders")
        tcb.append(_border_el("bottom", width))
        tcPr.append(tcb)


def set_cell_margins(table, vertical=14, horizontal=80):
    """Table-wide cell padding in twips (1/20 pt).

    Word defaults to 0 vertical / 108 horizontal which, combined with the Normal
    style's paragraph spacing, makes rows far taller than a typeset table.
    These values plus the zero paragraph spacing set in `_write_cell` are the
    density a journal prints at."""
    tblPr = table._tbl.tblPr
    for el in tblPr.findall(qn("w:tblCellMar")):
        tblPr.remove(el)
    mar = OxmlElement("w:tblCellMar")
    for tag, val in (("top", vertical), ("left", horizontal),
                     ("bottom", vertical), ("right", horizontal)):
        el = OxmlElement("w:%s" % tag)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)


def repeat_header_row(row):
    """Reprint this row atop each continuation page.

    Without it, a table that spills over opens the next page with bare numbers
    and no column labels."""
    trPr = row._tr.get_or_add_trPr()
    for el in trPr.findall(qn("w:tblHeader")):
        trPr.remove(el)
    trPr.append(OxmlElement("w:tblHeader"))


def _set_col_widths(table, col_widths):
    """Fixed layout so Word honours per-column widths instead of auto-fitting.

    Widths must be stamped on every cell, not just `table.columns[i].width`, and
    on `tblGrid` as well. `add_table` fills tblGrid with equal columns, and that
    stale grid outlives the per-cell widths: Word honours the cells under a fixed
    layout, but LibreOffice, journal conversion tools and any audit that reads the
    grid see the original equal split and report a table that is not the one Word
    draws."""
    table.autofit = False
    table.allow_autofit = False
    tblPr = table._tbl.tblPr
    for el in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(el)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for i, w in enumerate(col_widths):
        for row in table.rows:
            if i < len(row.cells):
                row.cells[i].width = Inches(w)
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for i, gc in enumerate(grid.findall(qn("w:gridCol"))):
            if i < len(col_widths):
                gc.set(qn("w:w"), str(int(round(col_widths[i] * 1440))))


def _write_cell(p, text, font_size, bold=False, font=DEFAULT_FONT):
    """Write cell text, honouring embedded newlines, at tight spacing.

    A raw "\\n" in a run is ignored by Word, so multi-line headers such as
    "Overall\\n(N=19,959)" need explicit breaks."""
    lines = str(text).split("\n")
    for k, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = font
        run.font.size = Pt(font_size)
        run.bold = bold
        if k < len(lines) - 1:
            run.add_break()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    # Size the paragraph mark too. An empty cell has no sized run, so its line
    # height falls back to the Normal style (12pt) and that row alone grows -
    # which is why rows with a blank p-value used to sit taller than the rest.
    pPr = p._p.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.insert(0, rPr)
    for tag in ("w:sz", "w:szCs"):
        for el in rPr.findall(qn(tag)):
            rPr.remove(el)
        el = OxmlElement(tag)
        el.set(qn("w:val"), str(int(font_size * 2)))     # half-points
        rPr.append(el)
    return p


# ---------------------------------------------------------------------------
# the one function most callers need
# ---------------------------------------------------------------------------

def add_journal_table(doc, rows, col_widths=None, header_rows=1, font_size=9,
                      font=DEFAULT_FONT, indent=0.12, group_space_before=0):
    """Append `rows` to `doc` as a journal-style table and return it.

    rows         list of lists of strings; rows[0] is the header
    col_widths   inches per column. Always pass these: without them Word
                 auto-fits and short columns (P, N) take room the labels need
    header_rows  number of leading header rows (bolded, ruled under, repeated)
    indent       paragraph indent applied to rows whose label starts with two
                 spaces (Word collapses the spaces themselves)

    Group header rows - body rows with no data except possibly a p-value - are
    bolded, merged across the width so a long label cannot wrap inside a narrow
    column, and held to the following row so a label never strands at the foot
    of a page. They take the same height as a data row: bold alone separates the
    variables, and an extra leading pushed the rows out of a single rhythm."""
    if not rows:
        return None
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_journal_borders(table)
    set_cell_margins(table)

    # the p-value column may be filled on a group row, so it neither disqualifies
    # the row as a group header nor gets swallowed by the merge
    pcols = pvalue_columns(rows[0]) if header_rows else set()
    merge_to = ncols - 1
    while merge_to > 0 and merge_to in pcols:
        merge_to -= 1

    group_rows = []
    for i, row in enumerate(rows):
        body = i >= header_rows
        group = body and is_group_row(row, ignore_cols=pcols)
        if group:
            group_rows.append(i)
        for j in range(ncols):
            text = str(row[j]) if j < len(row) else ""
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            if j == 0 and body and text.startswith("  "):
                text = text.lstrip()
                p.paragraph_format.left_indent = Inches(indent)
            # bold the column header and the variable label, but not a p-value
            # riding on a group row - that is ordinary data
            _write_cell(p, text, font_size, bold=(i < header_rows or (group and j == 0)),
                        font=font)
            if group and j == 0:
                p.paragraph_format.space_before = Pt(group_space_before)
                p.paragraph_format.keep_with_next = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if header_rows:
        set_row_bottom_rule(table.rows[header_rows - 1])
        for i in range(header_rows):
            repeat_header_row(table.rows[i])
    if col_widths:
        _set_col_widths(table, col_widths)

    # merge last: merging rewrites a row's cell list, so widths must be set first
    for i in group_rows:
        if merge_to <= 0:
            continue
        cells = table.rows[i].cells
        merged = cells[0].merge(cells[merge_to])
        # merge() concatenates the paragraphs of every cell it absorbs, so the
        # merged cell inherits one empty paragraph per column and the row
        # balloons to several lines tall. Keep only the label's paragraph.
        for p in merged.paragraphs[1:]:
            p._element.getparent().remove(p._element)
    return table
