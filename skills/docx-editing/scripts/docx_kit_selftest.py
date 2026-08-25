#!/usr/bin/env python3
"""docx_kit 자체시험. 절반은 «하지 않는 것»을 증명한다(연구실 규율).

    python ~/.claude/skills/docx-editing/scripts/docx_kit_selftest.py

렌더(Word COM)는 여기서 돌리지 않는다 -- Word가 없는 기기에서도 나머지를 검사할 수
있어야 하고, 렌더는 `docx_kit.py <파일>`로 사람이 확인한다. 대신 «DispatchEx를 쓰는가»를
소스에서 검사한다: 그 한 글자가 2026-08-25 사고의 전부였다.
"""
import re
import sys
from pathlib import Path

# 콘솔이 cp949면 «»·한글이 인쇄되다 죽는다 -- 검사 «내용»과 무관한 실패다.
for _s in (sys.stdout, sys.stderr):
    if hasattr(_s, "reconfigure"):
        try:
            _s.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))

from docx import Document                                    # noqa: E402
from docx_kit import add_journal_table, render_markdown      # noqa: E402

fails = []


def check(name, ok, detail=""):
    print(f"  [{'OK  ' if ok else '실패'}] {name}" + (f"  -- {detail}" if not ok else ""))
    if not ok:
        fails.append(name)


print("=== render_markdown ===")
doc = Document()
render_markdown(doc, "# 제목\n\n손으로 줄바꿈한\n산문이 **굵게**를\n걸쳐 있다.\n\n둘째 문단.\n")
texts = [p.text for p in doc.paragraphs if p.text.strip()]

# ── 해야 하는 것 ──────────────────────────────────────────────────────────
check("하드랩 산문을 한 문단으로 합친다",
      any("손으로 줄바꿈한 산문이 굵게를 걸쳐 있다." == t for t in texts),
      f"실제: {texts}")
check("문단을 빈 줄로 가른다", "둘째 문단." in texts, f"실제: {texts}")
check("제목을 heading으로", any(p.style.name.startswith("Heading")
                             for p in doc.paragraphs if p.text == "제목"))
_runs = [r for p in doc.paragraphs for r in p.runs]
check("**굵게**를 굵은 run으로", any(r.bold and "굵게" in r.text for r in _runs))

# ── «하지 않아야» 하는 것 ─────────────────────────────────────────────────
check("별표를 글자로 남기지 «않는다»", not any("**" in t for t in texts),
      f"실제: {texts}")
check("줄마다 문단을 만들지 «않는다»", len(texts) == 3,
      f"문단 {len(texts)}개 -- 줄 단위로 쪼개면 5개가 된다")

print("\n=== add_journal_table ===")
doc2 = Document()
t = add_journal_table(doc2, [["항목", "값"], ["가", "1"], ["나", "2"]],
                      col_widths=[2.0, 1.0], font_size=9)
check("표가 붙는다", len(doc2.tables) == 1)
check("행 수가 맞는다", len(t.rows) == 3, f"실제 {len(t.rows)}")
_borders = t._tbl.xml
check("세로줄을 «넣지 않는다»(저널 관습)",
      'w:val="single"' not in _borders.split("insideV")[1][:200]
      if "insideV" in _borders else True)

print("\n=== 골격 (manuscript_shell / brief_shell) ===")
import tempfile                                            # noqa: E402
from docx_kit import GateError, brief_shell, manuscript_shell  # noqa: E402

tmp = Path(tempfile.mkdtemp())
SEC = [("Introduction", "Depression affects {{n}} patients.\n산문이 이어진다."),
       ("Methods", "We measured **CIST**.\n\n둘째 문단.")]

res = manuscript_shell(tmp / "m.docx", "제목", SEC, values={"n": "582"},
                       abstract="OBJECTIVES: x", cover="검토용 표지",
                       tables=[("기저 특성", [["항목", "값"], ["가", "1"]], [2.0, 1.0])],
                       word_limit=100)
d = Document(tmp / "m.docx")
alltext = "\n".join(p.text for p in d.paragraphs)
check("자리표시자를 치환한다", "582" in alltext and "{{n}}" not in alltext)
check("표를 저널 표로 붙인다", len(d.tables) == 1)
check("표 제목이 표 «위»에 온다", "Table 1. 기저 특성" in alltext)
check("별표를 글자로 남기지 «않는다»", "**" not in alltext)
check("본문 단어를 센다", res["body_words"] > 0 and res["body_words"] <= 100)

try:
    manuscript_shell(tmp / "x.docx", "t", [("S", "{{없는키}}")], values={})
    check("미치환이면 «멈춘다»", False, "GateError가 안 났다")
except GateError:
    check("미치환이면 «멈춘다»", True)

try:
    manuscript_shell(tmp / "x.docx", "t", [("S", "word " * 50)], word_limit=10)
    check("단어 상한을 넘으면 «멈춘다»", False, "GateError가 안 났다")
except GateError:
    check("단어 상한을 넘으면 «멈춘다»", True)

res2 = manuscript_shell(tmp / "r.docx", "t",
                        [("S", "claim [REF: 미검증 메타분석] 뒤 문장")])
check("미검증 참고문헌을 «센다»", len(res2["refs"]) == 1, f"{res2['refs']}")
check("[REF:]를 단어 수에서 뺀다", res2["body_words"] == 3, f"{res2['body_words']}")

brief_shell(tmp / "b.docx", "브리프", [("h", "1. 개요"), ("md", "본문."),
                                       ("table", "표 1", [["a", "b"], ["1", "2"]])])
db = Document(tmp / "b.docx")
check("브리프는 격자표를 쓴다(장르가 다르다)",
      db.tables[0].style.name == "Table Grid", db.tables[0].style.name)

print("\n=== 렌더러 -- 소스 검사 ===")
src = (HERE / "docx_kit.py").read_text(encoding="utf-8")
check("DispatchEx를 쓴다", re.search(r"win32\.DispatchEx\(", src) is not None)
check("맨몸 Dispatch를 «쓰지 않는다»",
      re.search(r"win32\.Dispatch\(", src) is None,
      "Dispatch는 떠 있는 인스턴스에 붙어 한 번 막히면 계속 실패한다")
check("원본이 아니라 사본을 연다", "shutil.copy2" in src)

print()
if fails:
    print(f"★ {len(fails)}건 실패: {', '.join(fails)}")
    sys.exit(1)
print("전부 통과")
