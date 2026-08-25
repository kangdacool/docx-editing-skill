#!/usr/bin/env python
"""md_to_docx.py -- render a plain markdown document to a readable .docx.

Built for the data-request paperwork, which has to be read and edited by a person
and pasted into web forms, not typeset for a journal. Deliberately small: headings,
paragraphs, bold/italic runs, bullet and numbered lists, block quotes and pipe
tables. No figures, no styles beyond the built-ins.

    python tools/md_to_docx.py <in.md> [out.docx]
"""
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")


def add_runs(par, text):
    """Render **bold**, `code` and *italic* inside one paragraph."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            par.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = par.add_run(piece[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            par.add_run(piece[1:-1]).italic = True
        else:
            par.add_run(piece)


def unwrap(lines):
    """Join hard-wrapped continuation lines back onto their list item or quote.

    A bold span that crosses a line break inside a list item -- '**MCPS\\n
    publication**' -- otherwise ends up split across two paragraphs and the
    asterisks are rendered literally. Found exactly that way in the MCPS
    registration draft, so the joining happens before parsing rather than being
    patched per-case afterwards.
    """
    out = []
    for raw in lines:
        line = raw.rstrip()
        prev = out[-1] if out else ""
        # two spaces is a valid list continuation in markdown, not three -- requiring
        # three silently split '**bold\n  text**' across paragraphs and leaked the
        # asterisks into the rendered document.
        is_cont = (line.startswith(("  ", "\t")) and line.strip()
                   and not line.strip().startswith(("|", ">", "-", "*", "#", "```"))
                   and not re.match(r"^\s*\d+\.\s", line))
        parent_is_item = bool(re.match(r"^(\s*)([-*]|\d+\.)\s+", prev)) or \
            prev.lstrip().startswith(">")
        if is_cont and parent_is_item:
            out[-1] = prev + " " + line.strip()
        else:
            out.append(line)
    return out


def flush_table(doc, rows):
    rows = [r for r in rows if not re.fullmatch(r"\s*\|[\s:|-]+\|\s*", r)]
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    if not cells:
        return
    t = doc.add_table(rows=len(cells), cols=max(len(c) for c in cells))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(cells):
        for j, val in enumerate(row):
            par = t.cell(i, j).paragraphs[0]
            add_runs(par, val)
            if i == 0:
                for r in par.runs:
                    r.bold = True
    doc.add_paragraph()


def render_markdown(doc, text):
    """마크다운 «본문»을 주어진 doc에 이어 붙인다. 문단·굵게·목록·인용·표·코드.

    ⭐⭐ 2026-08-25에 main()에서 뺐다. 이 루프가 main() 안에 갇혀 있어서 다른
       빌더가 재사용하지 못하고 «더 못한 판»을 다시 짰다(굵게만 처리하고 하드랩
       문단 합치기를 새로 구현). 동작은 그대로다.

    ⚠️ 핵심은 buf/flush_para다 -- 88자에서 손으로 줄바꿈한 산문은 «빈 줄»까지
       모아 한 문단으로 합쳐야 한다. 줄마다 add_paragraph 하면 한 문장이 여러
       문단으로 쪼개지고, 괄호 짝 검사 같은 것이 전부 오탐이 된다.
    """
    lines = unwrap(text.splitlines())
    buf, table = [], []

    def flush_para():
        if buf:
            add_runs(doc.add_paragraph(), " ".join(buf))
            buf.clear()

    in_code = False
    for raw in lines:
        line = raw.rstrip()
        # fenced code block: keep the layout, drop the fences, no inline markup.
        # Without this the ``` lines and any * or ` inside them are rendered literally,
        # which is how a checkbox flow diagram ends up full of stray asterisks.
        if line.strip().startswith("```"):
            flush_para()
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Pt(14)
            r = p.add_run(raw if raw.strip() else " ")
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            continue
        if line.strip().startswith("|"):
            flush_para()
            table.append(line)
            continue
        if table:
            flush_table(doc, table)
            table = []
        if not line.strip():
            flush_para()
            continue
        if line.startswith("#"):
            flush_para()
            lvl = len(line) - len(line.lstrip("#"))
            doc.add_heading(line.lstrip("#").strip(), level=min(lvl, 4))
            continue
        if line.strip() in ("---", "***"):
            flush_para()
            continue
        if line.lstrip().startswith(">"):
            flush_para()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(22)
            add_runs(p, line.lstrip().lstrip(">").strip())
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            continue
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            flush_para()
            bullet = "List Bullet" if m.group(2) in "-*" else "List Number"
            p = doc.add_paragraph(style=bullet)
            add_runs(p, m.group(3))
            continue
        buf.append(line.strip())

    flush_para()
    if table:
        flush_table(doc, table)

    return doc


def main():
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(src)[0] + ".docx"
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(7)

    render_markdown(doc, open(src, encoding="utf-8").read())

    doc.save(out)
    print("wrote %s" % out)
    print("  paragraphs: %d   tables: %d" % (len(doc.paragraphs), len(doc.tables)))


if __name__ == "__main__":
    main()
